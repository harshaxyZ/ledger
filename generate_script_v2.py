import docx
import os

doc = docx.Document()
doc.add_heading('Ledger Presentation Script', 0)

script_data = [
    ("Presenter 1 (Slides 2-4)", [
        ("Slide 2: Project Details", "Welcome everyone. We are here to present Ledger, our offline-first, privacy-focused expense tracker. Let's get started."),
        ("Slide 3: Hero Slide", "Your Money. Your Rules. Ledger puts you in complete control. It's fully offline, instantly fast, and mathematically precise."),
        ("Slide 4: Why Ledger Exists", "We built this because modern financial apps are fundamentally broken. They are bloated with ads, require constant internet, and actively harvest your personal spending data.")
    ]),
    ("Presenter 2 (Slides 5-7)", [
        ("Slide 5: The Problem", "The core problem is trust and speed. When you rely on third-party servers, your data is exposed. When you rely on cloud syncing, tracking a simple coffee purchase takes too long."),
        ("Slide 6: Objectives", "Our objectives were absolute: 100% offline capability, zero-latency transaction logging, and strict privacy by design. We wanted an app that feels like a native OS feature."),
        ("Slide 7: Literature Survey", "Looking at existing tools: Splitwise forces you onto their servers. Mint harvested data before shutting down. Wallet locks basic features behind premium subscriptions. Ledger completely removes these pain points.")
    ]),
    ("Presenter 3 (Slides 8-10)", [
        ("Slide 8: Existing vs Ledger", "Unlike legacy systems that push your data to the cloud to serve you ads, Ledger keeps everything strictly local. No accounts, no emails, no tracking. Total anonymity."),
        ("Slide 9: What Makes Us Different", "What truly separates Ledger is its sub-second performance. Because there are no centralized databases, opening the app and logging an expense is instantaneous. We've also embedded an AI coach that runs locally."),
        ("Slide 10: How Ledger Works", "The workflow is frictionless. You open the app, log your transaction, and immediately close it. All complex data aggregation happens locally on your device in the background.")
    ]),
    ("Presenter 4 (Slides 11-13)", [
        ("Slide 11: Architecture", "Our architecture is fundamentally decentralized. The React frontend interacts with custom state hooks, which write directly to the browser's LocalStorage. Our UI then instantly re-renders without any network calls."),
        ("Slide 12: Technology Stack", "We achieved this using React 19 and Vite. We styled the interface with Tailwind CSS for that premium dark-mode aesthetic, and utilized the native LocalStorage API for our database."),
        ("Slide 13: Core Features", "Ledger's core features include deep categorized logging, habit-forming gamification streaks to keep you consistent, and an AI Advisor that acts as your personal financial coach.")
    ]),
    ("Presenter 5 (Slides 14-16)", [
        ("Slide 14: Budget Warnings", "As you can see in this screenshot, our Budget system is dynamic. If you approach your limit, the interface reacts in real-time, flashing crimson red to instantly warn you about overspending."),
        ("Slide 15: Analytics & Insights", "Our Insights dashboard visualizes your spending trends. It instantly aggregates your top categories and monthly progress, presenting complex data in clean, digestible visual progress bars."),
        ("Slide 16: Privacy First Design", "Because there is no backend, we couldn't track you even if we wanted to. Your financial footprint remains entirely under your control, giving you true data ownership.")
    ]),
    ("Presenter 6 (Slides 17-23)", [
        ("Slide 17: Offline First Experience", "Thanks to our Service Workers, Ledger boots up instantly whether you are on 5G or completely off the grid. It is a true Progressive Web App."),
        ("Slide 18: Implementation Modules", "We divided development into core modules: The Transaction Engine for speed, the Budget System for limits, and the AI Coach integration via Groq for intelligence."),
        ("Slide 19: Results", "The results speak for themselves. We achieved 0ms logging latency, a native application feel, and successfully proved that high-end software can be 100% private."),
        ("Slide 20: Future Enhancements", "Our future roadmap includes custom user-defined categories, automated recurring logs, and a biometric lock for an extra layer of local device security."),
        ("Slide 21: Conclusion", "In conclusion, Ledger proves that powerful financial tracking doesn't require sacrificing your privacy. It is fast, secure, and fully scalable."),
        ("Slide 22: References", "Our development relied heavily on the official React Documentation, MDN's PWA specifications, and the Tailwind CSS design guidelines."),
        ("Slide 23: Thank You", "Thank you for your time and attention. Ledger is private by design. We are now open for any questions.")
    ])
]

for presenter, lines in script_data:
    doc.add_heading(presenter, level=1)
    for title, text in lines:
        p = doc.add_paragraph()
        p.add_run(title + ": ").bold = True
        p.add_run(text)

doc_path = os.path.abspath('Ledger_Presentation_Script_V2.docx')
doc.save(doc_path)
print(f"Saved DOCX to {doc_path}")
